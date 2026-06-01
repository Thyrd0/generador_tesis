from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT
from reportlab.lib import colors
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from datetime import datetime
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Intentar registrar Arial Narrow o Arial desde el sistema Windows
def register_arial_narrow():
    windir = os.environ.get("WINDIR", "C:\\Windows")
    font_paths = [
        (windir + "\\Fonts\\arialn.ttf", windir + "\\Fonts\\arialnb.ttf", "Arial-Narrow"),
        (windir + "\\Fonts\\arial.ttf", windir + "\\Fonts\\arialbd.ttf", "Arial-Narrow"),  # Fallback a Arial normal
    ]
    
    for reg_path, bold_path, name in font_paths:
        if os.path.exists(reg_path):
            try:
                pdfmetrics.registerFont(TTFont(name, reg_path))
                if os.path.exists(bold_path):
                    pdfmetrics.registerFont(TTFont(name + "-Bold", bold_path))
                else:
                    pdfmetrics.registerFont(TTFont(name + "-Bold", reg_path))
                return True
            except Exception as e:
                print(f"Error registrando fuente {reg_path}: {e}")
    return False

# Registrar fuente al importar
ARIAL_NARROW_AVAILABLE = register_arial_narrow()


class NumberedCanvas(canvas.Canvas):
    """
    Canvas personalizado para ReportLab que permite calcular el número de páginas total
    y colocar la numeración arábiga en la esquina inferior derecha a partir de la página 2.
    """
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_number(num_pages)
            super().showPage()
        super().save()

    def draw_page_number(self, page_count):
        # Omitir numeración en la carátula (página 1)
        if self._pageNumber == 1:
            return
            
        self.saveState()
        font_name = "Arial-Narrow" if ARIAL_NARROW_AVAILABLE else "Helvetica"
        self.setFont(font_name, 10)
        
        # Texto de numeración arábiga simple
        page_text = f"{self._pageNumber}"
        
        width, height = self._pagesize
        # X: margen derecho es de 2.5 cm, por lo que imprimimos en (ancho - 2.5 cm)
        x = width - 2.5*cm
        # Y: margen inferior es de 2.5 cm, imprimimos a 1.2 cm para que quede bien posicionado
        y = 1.2*cm
        
        self.drawRightString(x, y, page_text)
        self.restoreState()


class DocumentGenerator:
    def __init__(self):
        self.output_dir = "generated_docs"
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)

    def generar_pdf(self, contenido, request):
        # Aseguramos que la carpeta exista dentro de backend si aplica
        filename = f"{self.output_dir}/tesis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        
        # Configuración del documento con márgenes exactos: izquierdo 3 cm, el resto 2.5 cm
        doc = SimpleDocTemplate(
            filename,
            pagesize=letter,
            rightMargin=2.5*cm,
            leftMargin=3*cm,
            topMargin=2.5*cm,
            bottomMargin=2.5*cm
        )
        
        styles = getSampleStyleSheet()
        font_family = "Arial-Narrow" if ARIAL_NARROW_AVAILABLE else "Helvetica"
        font_family_bold = "Arial-Narrow-Bold" if ARIAL_NARROW_AVAILABLE else "Helvetica-Bold"
        
        # Estilos personalizados (Interlineado 1.5 a 12pt = leading 18pt)
        styles.add(ParagraphStyle(
            name='Justify',
            parent=styles['Normal'],
            alignment=TA_JUSTIFY,
            fontSize=12,
            fontName=font_family,
            leading=18
        ))
        
        styles.add(ParagraphStyle(
            name='Center',
            parent=styles['Normal'],
            alignment=TA_CENTER,
            fontSize=12,
            fontName=font_family,
            leading=18
        ))
        
        styles.add(ParagraphStyle(
            name='CenterBold',
            parent=styles['Normal'],
            alignment=TA_CENTER,
            fontSize=14,
            fontName=font_family_bold,
            leading=20
        ))
        
        styles.add(ParagraphStyle(
            name='Title',
            parent=styles['Title'],
            alignment=TA_CENTER,
            fontSize=16,
            fontName=font_family_bold,
            leading=24
        ))
        
        story = []
        
        # 1. Carátula (Página 1)
        story.append(Spacer(1, 0.5*inch))
        story.append(Paragraph("UNIVERSIDAD NACIONAL DE TRUJILLO", styles['CenterBold']))
        story.append(Paragraph("FACULTAD DE INGENIERÍA", styles['Center']))
        story.append(Paragraph("Programa de Estudios de Ingeniería de Sistemas", styles['Center']))
        story.append(Spacer(1, 1.2*inch))
        
        # Título en negrita y centrado
        story.append(Paragraph(contenido['caratula']['titulo'].upper(), styles['Title']))
        story.append(Spacer(1, 1.2*inch))
        
        # Autores y asesor
        autores_texto = " y ".join(contenido['caratula']['autores'])
        story.append(Paragraph(f"<b>AUTORES:</b><br/>{autores_texto.upper()}", styles['Center']))
        story.append(Spacer(1, 0.3*inch))
        story.append(Paragraph(f"<b>ASESOR:</b><br/>{contenido['caratula']['asesor'].upper()}", styles['Center']))
        story.append(Spacer(1, 0.3*inch))
        story.append(Paragraph(f"<b>LÍNEA DE INVESTIGACIÓN:</b><br/>{contenido['caratula']['linea_investigacion'].upper()}", styles['Center']))
        story.append(Spacer(1, 1*inch))
        story.append(Paragraph(f"{contenido['caratula']['ciudad'].upper()} - {contenido['caratula']['año']}", styles['Center']))
        
        story.append(PageBreak())
        
        # 2. Jurado Dictaminador (Página 2)
        story.append(Paragraph("JURADO DICTAMINADOR", styles['CenterBold']))
        story.append(Spacer(1, 0.4*inch))
        story.append(Paragraph("El presente proyecto de tesis ha sido evaluado y aprobado en la Universidad Nacional de Trujillo por el siguiente jurado dictaminador:", styles['Justify']))
        story.append(Spacer(1, 0.8*inch))
        
        style_jurado = ParagraphStyle(
            name='JuradoStyle',
            parent=styles['Normal'],
            fontName=font_family,
            fontSize=11,
            alignment=TA_CENTER,
            leading=16
        )
        
        vocal_nombre = contenido['jurado']['vocal']['nombre']
        jurado_table_data = [
            [
                Paragraph("<br/><br/><br/>_________________________________<br/><b>Dr. Juan Carlos Mendoza Ramirez</b><br/>Presidente", style_jurado),
                Paragraph("<br/><br/><br/>_________________________________<br/><b>Dr. Roberto Sanchez Gonzales</b><br/>Secretario", style_jurado)
            ],
            [
                Paragraph("<br/><br/><br/><br/>_________________________________<br/><b>Dr. " + vocal_nombre + "</b><br/>Vocal (Asesor)", style_jurado),
                ""
            ]
        ]
        
        jurado_table = Table(jurado_table_data, colWidths=[3.1*inch, 3.1*inch])
        jurado_table.setStyle(TableStyle([
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('BOTTOMPADDING', (0,0), (-1,-1), 15),
        ]))
        story.append(jurado_table)
        
        story.append(PageBreak())
        
        # 3. Índice General (Página 3)
        story.append(Paragraph("ÍNDICE GENERAL", styles['CenterBold']))
        story.append(Spacer(1, 0.4*inch))
        
        style_idx = ParagraphStyle(
            name='IndexItem',
            parent=styles['Normal'],
            fontName=font_family,
            fontSize=11,
            leading=16
        )
        
        # Simulamos un Índice General con alineación precisa con puntos guía
        toc_data = [
            [Paragraph("<b>JURADO DICTAMINADOR</b>", style_idx), Paragraph(". . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . .", style_idx), Paragraph("ii", style_idx)],
            [Paragraph("<b>ÍNDICE GENERAL</b>", style_idx), Paragraph(". . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . .", style_idx), Paragraph("iii", style_idx)],
            [Paragraph("<b>CAPÍTULO I: INTRODUCCIÓN</b>", style_idx), Paragraph(". . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . .", style_idx), Paragraph("1", style_idx)],
            [Paragraph("&nbsp;&nbsp;&nbsp;&nbsp;1.1. Realidad Problemática", style_idx), Paragraph(". . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . .", style_idx), Paragraph("1", style_idx)],
            [Paragraph("&nbsp;&nbsp;&nbsp;&nbsp;1.2. Antecedentes", style_idx), Paragraph(". . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . .", style_idx), Paragraph("3", style_idx)],
            [Paragraph("&nbsp;&nbsp;&nbsp;&nbsp;1.3. Marco Teórico", style_idx), Paragraph(". . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . .", style_idx), Paragraph("5", style_idx)],
            [Paragraph("&nbsp;&nbsp;&nbsp;&nbsp;1.4. Justificación de la Investigación", style_idx), Paragraph(". . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . .", style_idx), Paragraph("8", style_idx)],
            [Paragraph("&nbsp;&nbsp;&nbsp;&nbsp;1.5. Formulación del Problema", style_idx), Paragraph(". . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . .", style_idx), Paragraph("9", style_idx)],
            [Paragraph("&nbsp;&nbsp;&nbsp;&nbsp;1.6. Hipótesis de la Investigación", style_idx), Paragraph(". . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . .", style_idx), Paragraph("9", style_idx)],
            [Paragraph("&nbsp;&nbsp;&nbsp;&nbsp;1.7. Objetivos de la Investigación", style_idx), Paragraph(". . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . .", style_idx), Paragraph("10", style_idx)],
            [Paragraph("&nbsp;&nbsp;&nbsp;&nbsp;1.8. Limitaciones del Estudio", style_idx), Paragraph(". . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . .", style_idx), Paragraph("10", style_idx)],
            [Paragraph("<b>REFERENCIAS BIBLIOGRÁFICAS</b>", style_idx), Paragraph(". . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . .", style_idx), Paragraph("11", style_idx)],
            [Paragraph("<b>ANEXOS</b>", style_idx), Paragraph(". . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . .", style_idx), Paragraph("13", style_idx)],
            [Paragraph("&nbsp;&nbsp;&nbsp;&nbsp;Anexo 1: Árbol de Problemas", style_idx), Paragraph(". . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . .", style_idx), Paragraph("13", style_idx)],
            [Paragraph("&nbsp;&nbsp;&nbsp;&nbsp;Anexo 2: Árbol de Objetivos", style_idx), Paragraph(". . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . .", style_idx), Paragraph("14", style_idx)],
            [Paragraph("<b>DECLARACIÓN JURADA</b>", style_idx), Paragraph(". . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . .", style_idx), Paragraph("15", style_idx)]
        ]
        
        toc_table = Table(toc_data, colWidths=[2.2*inch, 3.6*inch, 0.4*inch])
        toc_table.setStyle(TableStyle([
            ('VALIGN', (0,0), (-1,-1), 'BOTTOM'),
            ('LEFTPADDING', (0,0), (-1,-1), 0),
            ('RIGHTPADDING', (0,0), (-1,-1), 0),
            ('BOTTOMPADDING', (0,0), (-1,-1), 2),
            ('TOPPADDING', (0,0), (-1,-1), 2),
        ]))
        story.append(toc_table)
        
        story.append(PageBreak())
        
        # 4. Introducción (Página 4+)
        story.append(Paragraph("CAPÍTULO I: INTRODUCCIÓN", styles['CenterBold']))
        story.append(Spacer(1, 0.3*inch))
        
        # Dividir por párrafos para aplicar justificación e interlineado 1.5 correctamente
        for parag in contenido['introduccion'].split("\n\n"):
            if parag.strip():
                story.append(Paragraph(parag.strip(), styles['Justify']))
                story.append(Spacer(1, 0.15*inch))
        
        story.append(PageBreak())
        
        # 5. Referencias Bibliográficas
        story.append(Paragraph("REFERENCIAS BIBLIOGRÁFICAS", styles['CenterBold']))
        story.append(Spacer(1, 0.3*inch))
        
        # Estilo APA con sangría francesa (de ser posible, o justificado con espacio)
        style_ref = ParagraphStyle(
            name='APAReference',
            parent=styles['Justify'],
            leftIndent=20,
            firstLineIndent=-20,
            fontSize=11,
            leading=16
        )
        
        for ref in contenido['referencias']:
            if ref.strip():
                story.append(Paragraph(ref.strip(), style_ref))
                story.append(Spacer(1, 0.15*inch))
        
        story.append(PageBreak())
        
        # 6. Anexos
        story.append(Paragraph("ANEXOS", styles['CenterBold']))
        story.append(Spacer(1, 0.4*inch))
        
        story.append(Paragraph("<b>ANEXO 1: ÁRBOL DE PROBLEMAS</b>", styles['Justify']))
        story.append(Spacer(1, 0.15*inch))
        
        # Estilo para los diagramas y textos de anexos (monospaciado o Arial chico)
        style_anexo_text = ParagraphStyle(
            name='AnexoText',
            parent=styles['Normal'],
            fontName=font_family,
            fontSize=10,
            leading=14
        )
        
        for p in contenido['anexos']['arbol_problemas'].split("\n"):
            story.append(Paragraph(p.replace(" ", "&nbsp;"), style_anexo_text))
            
        story.append(Spacer(1, 0.4*inch))
        story.append(Paragraph("<b>ANEXO 2: ÁRBOL DE OBJETIVOS</b>", styles['Justify']))
        story.append(Spacer(1, 0.15*inch))
        
        for p in contenido['anexos']['arbol_objetivos'].split("\n"):
            story.append(Paragraph(p.replace(" ", "&nbsp;"), style_anexo_text))
        
        # 7. Declaración Jurada
        story.append(PageBreak())
        story.append(Paragraph("DECLARACIÓN JURADA", styles['CenterBold']))
        story.append(Spacer(1, 0.4*inch))
        
        autores_nombres = ", ".join(contenido['caratula']['autores'])
        declaracion = f"""
        Yo(Nosotros), <b>{autores_nombres.upper()}</b>, identificado(s) con DNI N° __________________, 
        autor(es) del proyecto de tesis titulado <b>"{contenido['caratula']['titulo'].upper()}"</b>, 
        DECLARO(DECLARAMOS) BAJO JURAMENTO que:
        <br/><br/>
        1. El trabajo de investigación presentado es de mi (nuestra) autoría.
        <br/>
        2. Se han respetado estrictamente los derechos de autor, y todas las fuentes utilizadas están debidamente citadas de acuerdo con la norma APA v7.
        <br/>
        3. El trabajo no ha sido presentado anteriormente en ninguna institución para obtener grado académico o título profesional.
        <br/>
        4. Todos los datos presentados en el informe son verídicos y corresponden a la realidad de la investigación.
        <br/><br/>
        En señal de conformidad con los términos expuestos, firmo(amos) la presente declaración.
        <br/><br/><br/>
        {contenido['caratula']['ciudad']}, {datetime.now().strftime("%d de %B de %Y")}
        <br/><br/><br/><br/>
        __________________________________________<br/>
        <b>Firma del Autor(es)</b>
        """
        
        story.append(Paragraph(declaracion, styles['Justify']))
        
        # Construir PDF usando NumberedCanvas para la numeración arábiga en esquina inferior derecha
        doc.build(story, canvasmaker=NumberedCanvas)
        
        return filename

    def generar_docx(self, contenido, request):
        filename = f"{self.output_dir}/tesis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        doc = Document()
        
        # Configurar márgenes exactos: Superior/Inferior/Derecho 2.5 cm, Izquierdo 3 cm
        for section in doc.sections:
            section.top_margin = Inches(0.984)      # 2.5 cm
            section.bottom_margin = Inches(0.984)   # 2.5 cm
            section.right_margin = Inches(0.984)    # 2.5 cm
            section.left_margin = Inches(1.181)     # 3.0 cm
            
        # Estilo por defecto (Arial Narrow, 12pt, Interlineado 1.5, Justificado)
        style = doc.styles['Normal']
        style.font.name = 'Arial Narrow'
        style.font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # 1. Carátula
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = title_p.add_run("UNIVERSIDAD NACIONAL DE TRUJILLO\n")
        run.bold = True
        run.font.size = Pt(14)
        
        run = title_p.add_run("FACULTAD DE INGENIERÍA\n")
        run.font.size = Pt(12)
        
        run = title_p.add_run("Programa de Estudios de Ingeniería de Sistemas\n\n\n\n")
        run.font.size = Pt(12)
        
        run = title_p.add_run(contenido['caratula']['titulo'].upper() + "\n\n\n\n\n")
        run.bold = True
        run.font.size = Pt(16)
        
        autores_texto = " y ".join(contenido['caratula']['autores'])
        run = title_p.add_run(f"AUTORES:\n{autores_texto.upper()}\n\n")
        run.font.size = Pt(12)
        
        run = title_p.add_run(f"ASESOR:\n{contenido['caratula']['asesor'].upper()}\n\n")
        run.font.size = Pt(12)
        
        run = title_p.add_run(f"LÍNEA DE INVESTIGACIÓN:\n{contenido['caratula']['linea_investigacion'].upper()}\n\n\n\n")
        run.font.size = Pt(12)
        
        run = title_p.add_run(f"{contenido['caratula']['ciudad'].upper()} - {contenido['caratula']['año']}")
        run.font.size = Pt(12)
        
        doc.add_page_break()
        
        # 2. Jurado
        jurado_h = doc.add_paragraph()
        jurado_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = jurado_h.add_run("JURADO DICTAMINADOR")
        run.bold = True
        run.font.size = Pt(14)
        
        doc.add_paragraph("El presente proyecto de tesis ha sido evaluado y aprobado en la Universidad Nacional de Trujillo por el siguiente jurado dictaminador:\n\n")
        
        j_p = doc.add_paragraph()
        j_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        j_p.add_run("\n\n\n_________________________________\n")
        j_p.add_run("Dr. Juan Carlos Mendoza Ramirez\n").bold = True
        j_p.add_run("Presidente\n\n\n\n\n_________________________________\n")
        j_p.add_run("Dr. Roberto Sanchez Gonzales\n").bold = True
        j_p.add_run("Secretario\n\n\n\n\n_________________________________\n")
        j_p.add_run(f"Dr. {contenido['jurado']['vocal']['nombre']}\n").bold = True
        j_p.add_run("Vocal (Asesor)")
        
        doc.add_page_break()
        
        # 3. Índice General
        idx_h = doc.add_paragraph()
        idx_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = idx_h.add_run("ÍNDICE GENERAL")
        run.bold = True
        run.font.size = Pt(14)
        
        toc_items = [
            ("JURADO DICTAMINADOR", "ii"),
            ("ÍNDICE GENERAL", "iii"),
            ("CAPÍTULO I: INTRODUCCIÓN", "1"),
            ("    1.1. Realidad Problemática", "1"),
            ("    1.2. Antecedentes", "3"),
            ("    1.3. Marco Teórico", "5"),
            ("    1.4. Justificación de la Investigación", "8"),
            ("    1.5. Formulación del Problema", "9"),
            ("    1.6. Hipótesis de la Investigación", "9"),
            ("    1.7. Objetivos de la Investigación", "10"),
            ("    1.8. Limitaciones del Estudio", "10"),
            ("REFERENCIAS BIBLIOGRÁFICAS", "11"),
            ("ANEXOS", "13"),
            ("    Anexo 1: Árbol de Problemas", "13"),
            ("    Anexo 2: Árbol de Objetivos", "14"),
            ("DECLARACIÓN JURADA", "15")
        ]
        
        for item, page in toc_items:
            dots = "." * (80 - len(item) - len(page))
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r1 = p.add_run(item)
            if "CAPÍTULO" in item or "REFERENCIAS" in item or "ANEXOS" in item or "DECLARACIÓN" in item or "ÍNDICE" in item or "JURADO" in item:
                r1.bold = True
            p.add_run(dots + page)
            
        doc.add_page_break()
        
        # 4. Introducción
        intro_h = doc.add_paragraph()
        intro_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = intro_h.add_run("CAPÍTULO I: INTRODUCCIÓN")
        run.bold = True
        run.font.size = Pt(14)
        
        for parag in contenido['introduccion'].split("\n\n"):
            if parag.strip():
                doc.add_paragraph(parag.strip())
                
        doc.add_page_break()
        
        # 5. Referencias
        ref_h = doc.add_paragraph()
        ref_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = ref_h.add_run("REFERENCIAS BIBLIOGRÁFICAS")
        run.bold = True
        run.font.size = Pt(14)
        
        for ref in contenido['referencias']:
            if ref.strip():
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(-0.5)
                p.add_run(ref.strip())
                
        doc.add_page_break()
        
        # 6. Anexos
        anx_h = doc.add_paragraph()
        anx_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = anx_h.add_run("ANEXOS")
        run.bold = True
        run.font.size = Pt(14)
        
        p = doc.add_paragraph()
        p.add_run("ANEXO 1: ÁRBOL DE PROBLEMAS").bold = True
        
        p_prob = doc.add_paragraph()
        p_prob.paragraph_format.line_spacing = 1.15
        p_prob.add_run(contenido['anexos']['arbol_problemas'])
        
        p = doc.add_paragraph()
        p.add_run("\nANEXO 2: ÁRBOL DE OBJETIVOS").bold = True
        
        p_obj = doc.add_paragraph()
        p_obj.paragraph_format.line_spacing = 1.15
        p_obj.add_run(contenido['anexos']['arbol_objetivos'])
        
        doc.add_page_break()
        
        # 7. Declaración Jurada
        dj_h = doc.add_paragraph()
        dj_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = dj_h.add_run("DECLARACIÓN JURADA")
        run.bold = True
        run.font.size = Pt(14)
        
        autores_nombres = ", ".join(contenido['caratula']['autores'])
        dj_texto = f"""
Yo(Nosotros), {autores_nombres.upper()}, identificado(s) con DNI N° __________________, autor(es) del proyecto de tesis titulado "{contenido['caratula']['titulo'].upper()}", DECLARO(DECLARAMOS) BAJO JURAMENTO que:

1. El trabajo de investigación presentado es de mi (nuestra) autoría.
2. Se han respetado estrictamente los derechos de autor, y todas las fuentes utilizadas están debidamente citadas de acuerdo con la norma APA v7.
3. El trabajo no ha sido presentado anteriormente en ninguna institución para obtener grado académico o título profesional.
4. Todos los datos presentados en el informe son verídicos y corresponden a la realidad de la investigación.

En señal de conformidad con los términos expuestos, firmo(amos) la presente declaración.

{contenido['caratula']['ciudad']}, {datetime.now().strftime("%d de %B de %Y")}



__________________________________________
Firma del Autor(es)
"""
        doc.add_paragraph(dj_texto)
        
        doc.save(filename)
        return filename